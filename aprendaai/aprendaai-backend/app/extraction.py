import base64
import json
from pathlib import Path

import fitz  # PyMuPDF
import docx
from openai import OpenAI

from app.config import OPENAI_API_KEY, OPENAI_MODEL
from app.prompts import IMAGE_EXTRACTION_PROMPT


def _extract_pdf_pages_as_images(file_path: str, max_pages: int = 5) -> str:
    """Fallback: render PDF pages as images and use GPT-4o Vision to extract text."""
    client = OpenAI(api_key=OPENAI_API_KEY)
    doc = fitz.open(file_path)
    text_parts = []

    for page_num in range(min(len(doc), max_pages)):
        page = doc[page_num]
        # Render page as image (150 DPI for good quality without being too large)
        pix = page.get_pixmap(dpi=150)
        image_data = base64.b64encode(pix.tobytes("png")).decode("utf-8")

        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": IMAGE_EXTRACTION_PROMPT},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{image_data}"
                            },
                        },
                    ],
                }
            ],
            max_tokens=2000,
        )
        page_text = response.choices[0].message.content or ""
        if page_text.strip():
            text_parts.append(page_text.strip())

    doc.close()
    full_text = "\n\n".join(text_parts).strip()
    return full_text[:5000] if full_text else ""


def extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF, with Vision fallback for scanned PDFs."""
    doc = fitz.open(file_path)
    text_parts = []
    for page_num in range(min(len(doc), 10)):  # Max 10 pages
        page = doc[page_num]
        text_parts.append(page.get_text())
    doc.close()
    full_text = "\n".join(text_parts).strip()

    # If text extraction returned very little content, the PDF is likely scanned/image-based
    if len(full_text) < 50:
        return _extract_pdf_pages_as_images(file_path)

    return full_text[:5000]


def extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    full_text = "\n".join(text_parts).strip()
    return full_text[:5000] if full_text else ""


def extract_from_image(file_path: str) -> str:
    """Extract text from an image using GPT-4o Vision."""
    client = OpenAI(api_key=OPENAI_API_KEY)

    with open(file_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    ext = Path(file_path).suffix.lower()
    mime_type = "image/png" if ext == ".png" else "image/jpeg"

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": IMAGE_EXTRACTION_PROMPT},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{image_data}"
                        },
                    },
                ],
            }
        ],
        max_tokens=2000,
    )
    return response.choices[0].message.content or ""


def extract_content(file_path: str, input_type: str) -> str:
    """Extract text content from any supported file type."""
    if input_type == "pdf":
        return extract_from_pdf(file_path)
    elif input_type in ("doc", "docx"):
        return extract_from_docx(file_path)
    elif input_type in ("image", "jpg", "jpeg", "png"):
        return extract_from_image(file_path)
    else:
        return ""
